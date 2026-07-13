"""Geração de DOCX (Word) para Proposta e Contrato do CRM.

Implementação simples (sem conversão HTML->DOCX completa):
- Cabeçalho (título + metadados)
- Dados da Empresa / Cliente
- Tabela de itens (produtos/serviços)
- Conteúdo em texto (HTML strip)
- Assinaturas (informativas)
"""

from html.parser import HTMLParser
from io import BytesIO

from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.phone_utils import telefone_exibicao_brasileiro

from .emitente_documento import obter_dados_emitente_documento


class _HtmlToDocxBlocksParser(HTMLParser):
    """Parser simples para transformar HTML em blocos (parágrafos e listas).
    Suporta: <p>, <br>, <b>/<strong>, <i>/<em>, <ul>/<ol>/<li>.
    """

    def __init__(self):
        super().__init__()
        self.blocks = []  # list[dict]
        self._cur_runs = []
        self._cur_is_list_item = False
        self._in_list = False
        self._list_style = "List Bullet"
        self._bold = 0
        self._italic = 0

    def _flush_paragraph(self):
        runs = [r for r in self._cur_runs if (r.get("text") or "").strip() != ""]
        if runs:
            if self._cur_is_list_item:
                self.blocks.append({"type": "li", "runs": runs, "style": self._list_style})
            else:
                self.blocks.append({"type": "p", "runs": runs})
        self._cur_runs = []
        self._cur_is_list_item = False

    def _append_text(self, text: str):
        if not text:
            return
        self._cur_runs.append(
            {
                "text": text,
                "bold": self._bold > 0,
                "italic": self._italic > 0,
            },
        )

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in ("p",) or tag == "br":
            self._flush_paragraph()
        elif tag in ("strong", "b"):
            self._bold += 1
        elif tag in ("em", "i"):
            self._italic += 1
        elif tag in ("ul", "ol"):
            self._flush_paragraph()
            self._in_list = True
            self._list_style = "List Bullet" if tag == "ul" else "List Number"
        elif tag == "li":
            self._flush_paragraph()
            self._cur_is_list_item = self._in_list

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ("p",):
            self._flush_paragraph()
        elif tag in ("strong", "b"):
            self._bold = max(0, self._bold - 1)
        elif tag in ("em", "i"):
            self._italic = max(0, self._italic - 1)
        elif tag in ("ul", "ol"):
            self._flush_paragraph()
            self._in_list = False

    def handle_data(self, data):
        if not data:
            return
        txt = data.replace("\xa0", " ")
        self._append_text(txt)

    def finalize(self):
        self._flush_paragraph()
        # Normalização: junta múltiplos espaços dentro de cada run
        for b in self.blocks:
            for r in b.get("runs", []):
                r["text"] = " ".join((r.get("text") or "").split())
        self.blocks = [b for b in self.blocks if any((r.get("text") or "") for r in b.get("runs", []))]
        return self.blocks


def _add_html_as_docx(document, html: str):
    """Adiciona HTML ao DOCX com layout simples (parágrafos + listas).
    """
    if not html:
        document.add_paragraph("Conteúdo não informado.")
        return

    parser = _HtmlToDocxBlocksParser()
    parser.feed(str(html))
    blocks = parser.finalize()

    if not blocks:
        document.add_paragraph("Conteúdo não informado.")
        return

    for b in blocks:
        if b.get("type") == "li":
            p = document.add_paragraph(style=b.get("style") or "List Bullet")
        else:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run_data in b["runs"]:
            t = run_data.get("text") or ""
            if not t:
                continue
            run = p.add_run(t)
            run.bold = bool(run_data.get("bold"))
            run.italic = bool(run_data.get("italic"))


def _formatar_valor(valor) -> str:
    if valor is None:
        return "—"
    try:
        v = float(valor)
        return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except (TypeError, ValueError):
        return "—"


def _montar_endereco_loja(loja) -> str | None:
    """Monta endereço formatado da loja."""
    cidade = getattr(loja, "cidade", "") or ""
    uf = getattr(loja, "uf", "") or ""
    cidade_uf = f"{cidade}/{uf}" if (cidade and uf) else (cidade or uf)
    parts = [
        getattr(loja, "logradouro", "") or "",
        getattr(loja, "numero", "") or "",
        getattr(loja, "complemento", "") or "",
        getattr(loja, "bairro", "") or "",
        cidade_uf,
        f"CEP {loja.cep}" if getattr(loja, "cep", "") else "",
    ]
    return ", ".join(p for p in parts if p).strip() or None


def _montar_admin_loja(owner) -> tuple[str | None, str | None]:
    """Retorna (admin_nome, admin_email) a partir do owner da loja."""
    if not owner:
        return None, None
    admin_nome = f"{getattr(owner, 'first_name', '') or ''} {getattr(owner, 'last_name', '') or ''}".strip() or getattr(owner, "username", "") or None
    admin_email = getattr(owner, "email", None) or None
    return admin_nome, admin_email


def _obter_dados_loja(loja_id: int) -> dict:
    try:
        from superadmin.models import Loja
        loja = Loja.objects.using("default").filter(id=loja_id).select_related("owner").first()
        if not loja:
            return {}
        admin_nome, admin_email = _montar_admin_loja(getattr(loja, "owner", None))
        return {
            "nome": getattr(loja, "nome", "") or "",
            "endereco": _montar_endereco_loja(loja),
            "cpf_cnpj": getattr(loja, "cpf_cnpj", "") or None,
            "admin_nome": admin_nome,
            "admin_email": admin_email,
        }
    except Exception:
        return {}


def _formatar_endereco_lead(lead) -> str:
    if not lead:
        return "—"
    parts = [
        getattr(lead, "logradouro", "") or "",
        (getattr(lead, "numero", "") and f"nº {lead.numero}") or "",
        getattr(lead, "complemento", "") or "",
        getattr(lead, "bairro", "") or "",
        (
            f"{lead.cidade}/{lead.uf}"
            if (getattr(lead, "cidade", "") and getattr(lead, "uf", ""))
            else (getattr(lead, "cidade", "") or getattr(lead, "uf", ""))
        ),
        (getattr(lead, "cep", "") and f"CEP {lead.cep}") or "",
    ]
    return ", ".join(p for p in parts if p).strip() or "—"


def _formatar_nome_usuario(user) -> str:
    if not user:
        return "—"
    first_name = getattr(user, "first_name", "") or ""
    last_name = getattr(user, "last_name", "") or ""
    full = f"{first_name} {last_name}".strip()
    return full or getattr(user, "nome", "") or getattr(user, "username", "") or "—"


def _docx_add_dados_empresa(doc, loja: dict | None) -> None:
    """Adiciona seção de dados da empresa ao documento."""
    if not loja:
        return
    doc.add_heading("Dados da Empresa", level=2)
    doc.add_paragraph(f"Nome: {loja.get('nome') or '—'}")
    if loja.get("endereco"):
        doc.add_paragraph(f"Endereço: {loja.get('endereco')}")
    if loja.get("cpf_cnpj"):
        doc.add_paragraph(f"CPF/CNPJ: {loja.get('cpf_cnpj')}")
    if loja.get("admin_nome"):
        doc.add_paragraph(f"Administrador: {loja.get('admin_nome')}")
    if loja.get("admin_email"):
        doc.add_paragraph(f"Email do administrador: {loja.get('admin_email')}")


def _docx_add_dados_cliente(doc, lead) -> None:
    """Adiciona seção de dados do cliente ao documento."""
    if not lead:
        return
    doc.add_heading("Dados do Cliente", level=2)
    doc.add_paragraph(f"Nome: {getattr(lead, 'nome', '') or '—'}")
    if getattr(lead, "empresa", ""):
        doc.add_paragraph(f"Empresa: {lead.empresa}")
    if getattr(lead, "cpf_cnpj", ""):
        doc.add_paragraph(f"CPF/CNPJ: {lead.cpf_cnpj}")
    if getattr(lead, "email", ""):
        doc.add_paragraph(f"Email: {lead.email}")
    if getattr(lead, "telefone", ""):
        tel_fmt = telefone_exibicao_brasileiro(lead.telefone) or lead.telefone
        doc.add_paragraph(f"Telefone: {tel_fmt}")
    doc.add_paragraph(f"Endereço: {_formatar_endereco_lead(lead)}")


def _docx_add_tabela_itens(doc, itens: list) -> None:
    """Adiciona tabela de itens ao documento."""
    doc.add_heading("Produtos e Serviços da Oportunidade", level=2)
    if not itens:
        doc.add_paragraph("Nenhum item cadastrado.")
        return
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Qtd"
    hdr[2].text = "Preço Unit."
    hdr[3].text = "Subtotal"
    for item in itens:
        ps = getattr(item, "produto_servico", None)
        tipo_ps = ps.get_tipo_display() if ps and hasattr(ps, "get_tipo_display") else (getattr(ps, "tipo", "") if ps else "")
        nome = f"{tipo_ps}: {ps.nome}" if (ps and tipo_ps) else (getattr(ps, "nome", "—") if ps else "—")
        row = table.add_row().cells
        row[0].text = nome
        row[1].text = str(getattr(item, "quantidade", None) or 1)
        row[2].text = _formatar_valor(getattr(item, "preco_unitario", None))
        row[3].text = _formatar_valor(getattr(item, "subtotal", None))


def gerar_docx_proposta(proposta) -> BytesIO:
    from docx import Document

    doc = Document()
    doc.add_heading("PROPOSTA COMERCIAL", level=1)

    numero = getattr(proposta, "numero", None) or str(getattr(proposta, "id", ""))
    doc.add_paragraph(f"Número: {numero}")
    doc.add_paragraph(f"Título: {getattr(proposta, 'titulo', '') or '—'}")
    doc.add_paragraph(f"Valor total: {_formatar_valor(getattr(proposta, 'valor_total', None))}")

    loja = obter_dados_emitente_documento(proposta)
    _docx_add_dados_empresa(doc, loja)
    lead = (
        proposta.oportunidade.lead
        if getattr(proposta, "oportunidade", None) and getattr(proposta.oportunidade, "lead", None)
        else None
    )
    _docx_add_dados_cliente(doc, lead)
    itens = list(getattr(proposta.oportunidade, "itens", []).all()) if getattr(proposta, "oportunidade", None) else []
    _docx_add_tabela_itens(doc, itens)
    doc.add_heading("Conteúdo", level=2)
    _add_html_as_docx(doc, getattr(proposta, "conteudo", "") or "")
    doc.add_heading("Assinaturas", level=2)
    vendedor = (
        proposta.oportunidade.vendedor
        if getattr(proposta, "oportunidade", None) and getattr(proposta.oportunidade, "vendedor", None)
        else None
    )
    doc.add_paragraph(f"Vendedor: {_formatar_nome_usuario(vendedor)}")
    doc.add_paragraph(f"Cliente: {getattr(lead, 'nome', '') or '—'}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def gerar_docx_contrato(contrato) -> BytesIO:
    from docx import Document

    doc = Document()
    doc.add_heading("CONTRATO", level=1)

    numero = getattr(contrato, "numero", None) or str(getattr(contrato, "id", ""))
    doc.add_paragraph(f"Número: {numero}")
    doc.add_paragraph(f"Título: {getattr(contrato, 'titulo', '') or '—'}")
    doc.add_paragraph(f"Valor total: {_formatar_valor(getattr(contrato, 'valor_total', None))}")

    loja = obter_dados_emitente_documento(contrato)
    _docx_add_dados_empresa(doc, loja)
    lead = (
        contrato.oportunidade.lead
        if getattr(contrato, "oportunidade", None) and getattr(contrato.oportunidade, "lead", None)
        else None
    )
    _docx_add_dados_cliente(doc, lead)
    itens = list(getattr(contrato.oportunidade, "itens", []).all()) if getattr(contrato, "oportunidade", None) else []
    _docx_add_tabela_itens(doc, itens)
    doc.add_heading("Conteúdo", level=2)
    _add_html_as_docx(doc, getattr(contrato, "conteudo", "") or "")
    doc.add_heading("Assinaturas", level=2)
    vendedor = (
        contrato.oportunidade.vendedor
        if getattr(contrato, "oportunidade", None) and getattr(contrato.oportunidade, "vendedor", None)
        else None
    )
    doc.add_paragraph(f"Vendedor: {_formatar_nome_usuario(vendedor)}")
    doc.add_paragraph(f"Cliente: {getattr(lead, 'nome', '') or '—'}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

